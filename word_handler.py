"""Word document handling module for attendance tracking."""

import os
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, List
import logging
import calendar

from docx import Document
from docx.shared import Inches
from docx.table import Table, _Cell

logger = logging.getLogger(__name__)

class WordHandler:
    """Handles Word document reading, filling, and generation."""
    
    def __init__(self, config_manager):
        """Initialize Word handler.
        
        Args:
            config_manager: Configuration manager instance
        """
        self.config = config_manager
        self.last_filled_doc = None
        
    def fill_attendance_sheet(self, time_in: datetime = None, time_out: datetime = None) -> Optional[str]:
        """Fill attendance sheet with timestamps.
        
        Args:
            time_in: Login timestamp
            time_out: Logout timestamp (optional)
            
        Returns:
            Path to filled Word document or None on error
        """
        doc_path = self.config.get('document_path')
        
        # Check if we should use the last filled document for updates
        current_date = datetime.now().date()
        if (self.last_filled_doc and 
            os.path.exists(self.last_filled_doc) and 
            time_out and not time_in):  # This is a check-out operation
            
            # Check if the last document was created today
            last_doc_stat = os.path.getmtime(self.last_filled_doc)
            last_doc_date = datetime.fromtimestamp(last_doc_stat).date()
            
            if last_doc_date == current_date:
                logger.info(f"Using existing document for check-out: {self.last_filled_doc}")
                doc_path = self.last_filled_doc
        
        if not doc_path or not os.path.exists(doc_path):
            logger.error(f"Word document not found: {doc_path}")
            return None
        
        try:
            # Load the document
            doc = Document(doc_path)
            
            # Get current date and determine if it's a weekend
            is_weekend = self._is_weekend(current_date)
            
            # Add selected month to the document
            self._add_month_to_document(doc)
            
            # Find and fill the attendance table
            attendance_table = self._find_attendance_table(doc)
            if not attendance_table:
                logger.error("No attendance table found in document")
                return None
            
            # Fill the table based on current date and events
            self._fill_table_for_date(attendance_table, current_date, time_in, time_out, is_weekend)
            
            # If it's Friday and we're checking in, automatically fill weekend days
            if current_date.weekday() == 4 and time_in:  # Friday = 4
                self._fill_weekend_days(attendance_table, current_date)
            
            # Save the filled document
            if doc_path == self.last_filled_doc:
                # Update existing document
                output_path = doc_path
                logger.info(f"Updating existing document: {output_path}")
            else:
                # Create new document
                output_path = self._generate_output_path(doc_path)
                logger.info(f"Creating new document: {output_path}")
            
            doc.save(output_path)
            
            self.last_filled_doc = output_path
            logger.info(f"Word document filled successfully: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error filling Word document: {e}")
            return None
    
    def _is_weekend(self, date) -> bool:
        """Check if the given date is a weekend (Saturday or Sunday).
        
        Args:
            date: Date to check
            
        Returns:
            True if weekend, False otherwise
        """
        return date.weekday() >= 5  # Saturday = 5, Sunday = 6
    
    def _add_month_to_document(self, doc: Document):
        """Add selected month to the document.
        
        Args:
            doc: Word document to modify
        """
        selected_month = self.config.get('selected_month', '')
        if not selected_month:
            return
        
        # Look for paragraphs that contain "Month:" and fill the value after it
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if 'Month:' in text and not selected_month in text:
                # Replace the text after "Month:" with the selected month
                paragraph.text = f"Month: {selected_month}"
                logger.info(f"Added month '{selected_month}' to document")
                return
        
        logger.warning("No 'Month:' field found in document")
    
    def _find_attendance_table(self, doc: Document) -> Optional[Table]:
        """Find the attendance table in the document.
        
        Args:
            doc: Word document
            
        Returns:
            Table object or None if not found
        """
        # Look for tables that might contain attendance data
        for table in doc.tables:
            if self._is_attendance_table(table):
                return table
        return None
    
    def _is_attendance_table(self, table: Table) -> bool:
        """Check if a table is likely an attendance table.
        
        Args:
            table: Table to check
            
        Returns:
            True if likely an attendance table
        """
        # Check for common attendance table headers
        headers = []
        if table.rows:
            first_row = table.rows[0]
            for cell in first_row.cells:
                headers.append(cell.text.lower().strip())
        
        # Look for attendance-related keywords
        attendance_keywords = ['date', 'day', 'time in', 'time out', 'hours', 'attendance']
        matches = sum(1 for keyword in attendance_keywords if any(keyword in header for header in headers))
        
        return matches >= 2  # At least 2 keywords match
    
    def _fill_table_for_date(self, table: Table, current_date, time_in: datetime, time_out: datetime, is_weekend: bool):
        """Fill table for the current date.
        
        Args:
            table: Attendance table
            current_date: Current date
            time_in: Check-in time
            time_out: Check-out time
            is_weekend: Whether current date is weekend
        """
        # Find column indices
        header_row = table.rows[0]
        columns = self._identify_columns(header_row)
        
        if not columns:
            logger.warning("Could not identify table columns")
            return
        
        # Find or create row for current date
        target_row = self._find_or_create_date_row(table, current_date, columns)
        
        if target_row:
            self._fill_row_data(target_row, columns, current_date, time_in, time_out, is_weekend)
    
    def _identify_columns(self, header_row) -> Dict[str, int]:
        """Identify column indices based on headers.
        
        Args:
            header_row: Header row of the table
            
        Returns:
            Dictionary mapping column types to indices
        """
        columns = {}
        
        for i, cell in enumerate(header_row.cells):
            header_text = cell.text.lower().strip()
            logger.info(f"Column {i}: '{cell.text}' -> '{header_text}'")
            
            # Date column
            if any(keyword in header_text for keyword in ['date', 'dt']):
                columns['date'] = i
                logger.info(f"Found date column at index {i}")
            
            # Day column
            elif any(keyword in header_text for keyword in ['day', 'weekday']):
                columns['day'] = i
                logger.info(f"Found day column at index {i}")
            
            # Time in column - more specific matching
            elif 'time in' in header_text or header_text in ['time in', 'timein', 'in time', 'intime', 'check in', 'checkin', 'start', 'start time']:
                columns['time_in'] = i
                logger.info(f"Found time_in column at index {i}")
            
            # Time out column - more specific matching  
            elif 'time out' in header_text or header_text in ['time out', 'timeout', 'out time', 'outtime', 'check out', 'checkout', 'end', 'end time']:
                columns['time_out'] = i
                logger.info(f"Found time_out column at index {i}")
            
            # Hours column
            elif any(keyword in header_text for keyword in ['hours', 'total', 'duration']):
                columns['hours'] = i
                logger.info(f"Found hours column at index {i}")
        
        logger.info(f"Identified columns: {columns}")
        return columns
    
    def _find_or_create_date_row(self, table: Table, target_date, columns) -> Optional[Any]:
        """Find existing row for date or identify empty row to fill.
        
        Args:
            table: Attendance table
            target_date: Date to find/create row for
            columns: Column mapping
            
        Returns:
            Row to fill or None
        """
        date_col = columns.get('date')
        day_col = columns.get('day')
        
        # First, try to find a row that already has today's date
        for i, row in enumerate(table.rows[1:], 1):  # Skip header
            if i < len(table.rows):
                if date_col is not None:
                    date_cell = row.cells[date_col]
                    cell_text = date_cell.text.strip()
                    
                    # If cell contains target date, use this row
                    if self._is_date_match(cell_text, target_date):
                        logger.info(f"Found existing row for date: {target_date}")
                        return row
        
        # If no existing date found, find the first empty row or a row for today's day of week
        target_day_name = target_date.strftime('%A')  # Monday, Tuesday, etc.
        target_date_str = target_date.strftime('%d/%m/%Y')
        
        for i, row in enumerate(table.rows[1:], 1):  # Skip header
            if i < len(table.rows):
                # Check if this row is empty and can be used
                is_empty_row = True
                for j, cell in enumerate(row.cells):
                    if cell.text.strip() and cell.text.strip() not in ['-', '']:
                        is_empty_row = False
                        break
                
                if is_empty_row:
                    logger.info(f"Using empty row {i} for date: {target_date}")
                    return row
        
        logger.warning(f"No suitable row found for date: {target_date}")
        return None
    
    def _is_date_match(self, cell_text: str, target_date) -> bool:
        """Check if cell text matches target date.
        
        Args:
            cell_text: Text in the cell
            target_date: Date to match
            
        Returns:
            True if matches
        """
        if not cell_text:
            return False
        
        # Try different date formats
        date_formats = ['%Y-%m-%d', '%d/%m/%Y', '%m/%d/%Y', '%d-%m-%Y']
        
        for fmt in date_formats:
            try:
                parsed_date = datetime.strptime(cell_text.strip(), fmt).date()
                if parsed_date == target_date:
                    return True
            except ValueError:
                continue
        
        return False
    
    def _fill_row_data(self, row, columns: Dict[str, int], current_date, time_in: datetime, time_out: datetime, is_weekend: bool):
        """Fill row with attendance data.
        
        Args:
            row: Table row to fill
            columns: Column mapping
            current_date: Current date
            time_in: Check-in time
            time_out: Check-out time
            is_weekend: Whether it's a weekend
        """
        date_format = self.config.get('date_format', '%d/%m/%Y')
        time_format = self.config.get('time_format', '%H:%M')
        
        # Fill date column
        if 'date' in columns:
            date_cell = row.cells[columns['date']]
            if not date_cell.text.strip() or self._is_placeholder(date_cell.text):
                date_cell.text = current_date.strftime(date_format)
                logger.info(f"Filled date: {current_date.strftime(date_format)}")
        
        # Fill day column
        if 'day' in columns:
            day_cell = row.cells[columns['day']]
            if not day_cell.text.strip() or self._is_placeholder(day_cell.text):
                day_cell.text = calendar.day_name[current_date.weekday()]
                logger.info(f"Filled day: {calendar.day_name[current_date.weekday()]}")
        
        # Fill time in column
        if 'time_in' in columns:
            time_in_cell = row.cells[columns['time_in']]
            current_text = time_in_cell.text.strip()
            logger.info(f"Time In cell current content: '{current_text}'")
            
            if is_weekend:
                time_in_cell.text = "Weekend"  # Fill with "Weekend" for weekends
                logger.info("Filled time in with 'Weekend' - weekend")
            elif time_in:
                # Fill time in if cell is empty or placeholder
                if not current_text or self._is_placeholder(current_text):
                    time_in_cell.text = time_in.strftime(time_format)
                    logger.info(f"Filled time in: {time_in.strftime(time_format)}")
                else:
                    logger.info(f"Time in already filled with: {current_text}")
            else:
                # If no time_in provided, keep existing value
                logger.info(f"No time_in provided - keeping existing value: '{current_text}'")
        else:
            logger.warning("time_in column not found in table")
        
        # Fill time out column
        if 'time_out' in columns:
            time_out_cell = row.cells[columns['time_out']]
            current_text = time_out_cell.text.strip()
            logger.info(f"Time Out cell current content: '{current_text}'")
            
            if is_weekend:
                time_out_cell.text = "Weekend"  # Fill with "Weekend" for weekends
                logger.info("Filled time out with 'Weekend' - weekend")
            elif time_out:
                # Fill time out if cell is empty or placeholder
                if not current_text or self._is_placeholder(current_text):
                    time_out_cell.text = time_out.strftime(time_format)
                    logger.info(f"Filled time out: {time_out.strftime(time_format)}")
                else:
                    logger.info(f"Time out already filled with: {current_text}")
            else:
                # If no time_out provided, keep existing value
                logger.info(f"No time_out provided - keeping existing value: '{current_text}'")
        else:
            logger.warning("time_out column not found in table")
        
        # Calculate and fill hours if we can determine both times
        if 'hours' in columns and not is_weekend:
            hours_cell = row.cells[columns['hours']]
            
            # Try to get both times (either from parameters or existing cell values)
            actual_time_in = time_in
            actual_time_out = time_out
            
            # If we don't have time_in from parameter, try to parse it from the cell
            if not actual_time_in and 'time_in' in columns:
                time_in_text = row.cells[columns['time_in']].text.strip()
                if time_in_text and not self._is_placeholder(time_in_text):
                    try:
                        actual_time_in = datetime.strptime(time_in_text, time_format)
                        logger.info(f"Parsed existing time_in: {time_in_text}")
                    except ValueError:
                        logger.warning(f"Could not parse existing time_in: {time_in_text}")
            
            # If we don't have time_out from parameter, try to parse it from the cell  
            if not actual_time_out and 'time_out' in columns:
                time_out_text = row.cells[columns['time_out']].text.strip()
                if time_out_text and not self._is_placeholder(time_out_text):
                    try:
                        actual_time_out = datetime.strptime(time_out_text, time_format)
                        logger.info(f"Parsed existing time_out: {time_out_text}")
                    except ValueError:
                        logger.warning(f"Could not parse existing time_out: {time_out_text}")
            
            # Calculate hours if we have both times
            if actual_time_in and actual_time_out:
                if not hours_cell.text.strip() or self._is_placeholder(hours_cell.text):
                    duration = actual_time_out - actual_time_in
                    hours = duration.total_seconds() / 3600
                    hours_cell.text = f"{hours:.2f}"
                    logger.info(f"Calculated hours: {hours:.2f}")
            else:
                logger.info("Cannot calculate hours - missing time_in or time_out")
    
    def _is_placeholder(self, text: str) -> bool:
        """Check if text is a placeholder that should be replaced.
        
        Args:
            text: Text to check
            
        Returns:
            True if it's a placeholder
        """
        placeholders = ['', '-', '--', 'N/A', 'TBD', 'TIME', 'DATE', 'DAY', 'HH:MM', 'DD/MM/YYYY']
        return text.strip().upper() in [p.upper() for p in placeholders]
    
    def _generate_output_path(self, input_path: str) -> str:
        """Generate output path for filled document.
        
        Args:
            input_path: Original document path
            
        Returns:
            Output file path
        """
        input_file = Path(input_path)
        output_dir = Path(self.config.get('output_directory', 'filled_docs'))
        # Only create output_dir if/when a file is actually written

        if not output_dir.exists():
            output_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{input_file.stem}_filled_{timestamp}{input_file.suffix}"

        return str(output_dir / filename)
    
    def _fill_weekend_days(self, table: Table, friday_date):
        """Fill Saturday and Sunday rows automatically when checking in on Friday.
        
        Args:
            table: Attendance table
            friday_date: The Friday date
        """
        from datetime import timedelta
        from docx.shared import RGBColor
        
        logger.info("Auto-filling weekend days for Friday check-in")
        
        # Calculate Saturday and Sunday dates
        saturday_date = friday_date + timedelta(days=1)
        sunday_date = friday_date + timedelta(days=2)
        
        # Get column mapping
        header_row = table.rows[0]
        columns = self._identify_columns(header_row)
        
        if not columns:
            logger.warning("Could not identify columns for weekend filling")
            return
        
        # Fill Saturday
        saturday_row = self._find_next_empty_row(table, columns)
        if saturday_row:
            self._fill_weekend_row(saturday_row, columns, saturday_date, "Saturday")
            self._highlight_weekend_row(saturday_row)
            logger.info(f"Filled Saturday row: {saturday_date}")
        
        # Fill Sunday  
        sunday_row = self._find_next_empty_row(table, columns)
        if sunday_row:
            self._fill_weekend_row(sunday_row, columns, sunday_date, "Sunday")
            self._highlight_weekend_row(sunday_row)
            logger.info(f"Filled Sunday row: {sunday_date}")
    
    def _find_next_empty_row(self, table: Table, columns):
        """Find the next empty row in the table.
        
        Args:
            table: Attendance table
            columns: Column mapping
            
        Returns:
            Next empty row or None
        """
        for i, row in enumerate(table.rows[1:], 1):  # Skip header
            if i < len(table.rows):
                # Check if this row is empty and can be used
                is_empty_row = True
                for j, cell in enumerate(row.cells):
                    if cell.text.strip() and cell.text.strip() not in ['-', '', 'N/A']:
                        is_empty_row = False
                        break
                
                if is_empty_row:
                    logger.info(f"Found next empty row at index {i}")
                    return row
        
        logger.warning("No empty rows found for weekend filling")
        return None
    
    def _fill_weekend_row(self, row, columns: Dict[str, int], date, day_name: str):
        """Fill a weekend row with date and day, leaving time columns empty.
        
        Args:
            row: Table row to fill
            columns: Column mapping
            date: Weekend date
            day_name: Day name (Saturday/Sunday)
        """
        date_format = self.config.get('date_format', '%d/%m/%Y')
        
        # Fill date column
        if 'date' in columns:
            date_cell = row.cells[columns['date']]
            if not date_cell.text.strip() or self._is_placeholder(date_cell.text):
                date_cell.text = date.strftime(date_format)
        
        # Fill day column
        if 'day' in columns:
            day_cell = row.cells[columns['day']]
            if not day_cell.text.strip() or self._is_placeholder(day_cell.text):
                day_cell.text = day_name
        
        # Fill time columns with "Weekend" for weekend days
        if 'time_in' in columns:
            time_in_cell = row.cells[columns['time_in']]
            time_in_cell.text = "Weekend"
        
        if 'time_out' in columns:
            time_out_cell = row.cells[columns['time_out']]
            time_out_cell.text = "Weekend"
        
        if 'hours' in columns:
            hours_cell = row.cells[columns['hours']]
            hours_cell.text = ""
    
    def _highlight_weekend_row(self, row):
        """Highlight a weekend row with gray background.
        
        Args:
            row: Table row to highlight
        """
        try:
            from docx.shared import RGBColor
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            
            # Light gray color (RGB: 220, 220, 220)
            gray_color = "DCDCDC"
            
            for cell in row.cells:
                # Set cell background color
                cell_properties = cell._tc.get_or_add_tcPr()
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}"/>')
                cell_properties.append(shading)
                
            logger.info("Applied gray highlighting to weekend row")
            
        except Exception as e:
            logger.warning(f"Could not apply highlighting to weekend row: {e}")
    
    def get_available_templates(self) -> List[str]:
        """Get list of available Word templates.
        
        Returns:
            List of template file paths
        """
        templates = []
        for ext in ['.docx', '.doc']:
            templates.extend(Path('.').glob(f'*{ext}'))
        return [str(t) for t in templates]
